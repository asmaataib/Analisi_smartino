#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convertitore Markdown to Word
Converte il report tecnico da Markdown a formato Word (.docx)
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_markdown_to_word(md_file, docx_file):
    """Converte file Markdown in documento Word"""
    
    # Crea nuovo documento Word
    doc = Document()
    
    # Leggi il file Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    in_code_block = False
    
    for line in lines:
        line_stripped = line.strip()
        
        # Gestisci blocchi di codice
        if line_stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            # Aggiungi codice con font monospace
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Inches(0.1)
            continue
        
        # Gestisci titoli
        if line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:], level=1)
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith('#### '):
            doc.add_heading(line_stripped[5:], level=4)
        
        # Gestisci liste
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            doc.add_paragraph(line_stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line_stripped):
            # Lista numerata
            text = re.sub(r'^\d+\. ', '', line_stripped)
            doc.add_paragraph(text, style='List Number')
        
        # Gestisci separatori
        elif line_stripped == '---':
            doc.add_page_break()
        
        # Gestisci paragrafi normali
        elif line_stripped and not line_stripped.startswith('---'):
            # Rimuovi markdown formatting
            clean_line = line_stripped
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_line)  # Bold
            clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)      # Italic
            clean_line = re.sub(r'`(.*?)`', r'\1', clean_line)        # Code
            
            if clean_line:
                doc.add_paragraph(clean_line)
    
    # Salva il documento
    doc.save(docx_file)
    print(f'✅ Report convertito con successo in: {docx_file}')

if __name__ == '__main__':
    # Converti il report
    convert_markdown_to_word(
        'Technical_Report_Respiratory_Disease_ML_Analysis.md',
        'Technical_Report_Respiratory_Disease_ML_Analysis.docx'
    )
    print('🎯 Conversione completata!')